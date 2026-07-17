#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tutti Tschutti / R1 - Event-Abrechnung (Variante 3: voll automatisierte Pipeline)

Holt die Verkaufsdaten automatisch aus der SumUp-API, liest die TWINT/RaiseNow-CSV
ein und erstellt pro Event eine fertige Abrechnung als Excel (.xlsx) und PDF.

Berechnet pro Event:
  1) Gesamtumsatz (SumUp)        -> 13% Umsatzbeteiligung
  2) Hausgetraenke-Verbrauch     -> Menge x Einkaufspreis (Gratis/Barteam als CHF-0-Artikel
                                    sind automatisch dabei, weil sie als Artikel getippt werden)
  3) Zahlungs-Split (Abgleich)   -> Karte / TWINT / Bargeld

WICHTIG - KEIN DOPPELZAEHLEN:
  In eurem Workflow wird JEDE Bestellung in SumUp getippt, auch TWINT-Zahlungen
  (gebucht als 'Bargeld'). Der SumUp-Gesamtumsatz ist daher bereits vollstaendig.
  Die RaiseNow/TWINT-CSV wird NICHT zum Umsatz addiert, sondern nur fuer den
  Zahlungs-Abgleich verwendet:  Bargeld (echt) = SumUp-Cash - TWINT.

VORAUSSETZUNG:
  In SumUp muss jedes Getraenk ein Artikel im Katalog sein (nicht als freier Betrag
  tippen). Nur dann liefert die API die Einzelpositionen, die fuer die EK-Verrechnung
  noetig sind.

Aufruf (Beispiele):
  # Erst testen, ohne API/Keys - mit eingebauten Demodaten:
  python tutti_abrechnung.py --demo

  # Echtlauf fuer ein Event:
  python tutti_abrechnung.py \
      --event "Eröffnungsabend" \
      --von "2026-06-11T17:00:00+02:00" \
      --bis "2026-06-12T02:00:00+02:00" \
      --twint raisenow_export.csv \
      --config config.json \
      --ek ek_preise.csv

  # Welche payment_type / card_type Werte hat mein Konto? (zum Konfigurieren):
  python tutti_abrechnung.py --inspect --von ... --bis ... --config config.json
"""

import argparse
import csv
import datetime as dt
import json
import os
import sys
import unicodedata

# ----------------------------------------------------------------------------
# SumUp API-Endpunkte (zentral, falls SumUp die Versionen aendert -> hier anpassen).
# Pruefe im Zweifel: https://developer.sumup.com/api
# ----------------------------------------------------------------------------
API_BASE = "https://api.sumup.com"
EP_ME = API_BASE + "/v0.1/me"
EP_TX_HISTORY = API_BASE + "/v0.1/me/transactions/history"
EP_RECEIPT = API_BASE + "/v1.1/receipts/{tx_id}?mid={mid}"

# Wie die SumUp-Zahlungsarten in "Karte" vs "Bargeld(+TWINT)" einsortiert werden.
# Defaults sind ein erster Versuch. Mit --inspect siehst du die echten Werte deines
# Kontos und kannst sie hier bei Bedarf anpassen.
CARD_PAYMENT_TYPES = {"ECOM", "POS", "BOLETO", "RECURRING"}  # alles mit Karte/Reader
CASH_PAYMENT_TYPES = {"CASH"}                                # bar getippt (= bar + TWINT)


# ----------------------------------------------------------------------------
# Hilfsfunktionen
# ----------------------------------------------------------------------------
def norm(name):
    """Produktnamen vereinheitlichen (klein, ohne Akzente/Doppel-Leerzeichen).

    Mitarbeiter-Gratisartikel (0 CHF, Kategorie 'Mitarbeiter') matchen auf
    denselben EK wie der normale Artikel: Suffixe/Praefixe wie 'Mitarbeiter',
    '(MA)', 'Personal', 'Gratis' werden entfernt."""
    if name is None:
        return ""
    s = unicodedata.normalize("NFKD", str(name)).encode("ascii", "ignore").decode()
    s = " ".join(s.lower().split())
    for tag in ("mitarbeiter", "personal", "gratis", "(ma)", "ma:"):
        s = s.replace(tag, " ")
    return " ".join(s.replace("(", " ").replace(")", " ").split())


def parse_iso(s):
    """ISO-8601 String -> aware datetime."""
    s = s.strip().replace("Z", "+00:00")
    return dt.datetime.fromisoformat(s)


def load_config(path):
    if not path or not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_ek_prices(path):
    """CSV 'produkt,ek' -> dict {normierter_name: (anzeigename, ek)}."""
    ek = {}
    if not path or not os.path.exists(path):
        return ek
    with open(path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            if not row or len(row) < 2:
                continue
            name, raw = row[0].strip(), row[1].strip()
            if norm(name) in ("produkt", "name", "getraenk", "getränk"):
                continue  # Header
            try:
                ek[norm(name)] = (name, float(str(raw).replace(",", ".")))
            except ValueError:
                continue
    return ek


def load_twint_csv(path, von, bis, cfg):
    """RaiseNow/TWINT-Export einlesen, Betraege im Zeitfenster summieren.

    Spaltennamen sind je nach Export unterschiedlich. Standardmaessig wird nach
    einer Datums- und einer Betrags-Spalte gesucht; per config['twint'] ueberschreibbar.
    """
    if not path or not os.path.exists(path):
        return {"total": 0.0, "count": 0, "note": "keine TWINT-CSV angegeben"}

    tw = (cfg.get("twint") or {})
    amount_keys = [tw.get("amount_col")] if tw.get("amount_col") else \
        ["amount", "betrag", "gross", "brutto", "value", "total"]
    date_keys = [tw.get("date_col")] if tw.get("date_col") else \
        ["date", "datum", "created", "created_at", "timestamp", "zeit", "time"]

    total, count = 0.0, 0
    with open(path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        headers_norm = {norm(h): h for h in (reader.fieldnames or [])}
        amount_col = next((headers_norm[norm(k)] for k in amount_keys
                           if k and norm(k) in headers_norm), None)
        date_col = next((headers_norm[norm(k)] for k in date_keys
                         if k and norm(k) in headers_norm), None)
        if amount_col is None:
            return {"total": 0.0, "count": 0,
                    "note": "Betrags-Spalte nicht gefunden -> in config['twint']['amount_col'] setzen. "
                            "Gefundene Spalten: " + ", ".join(reader.fieldnames or [])}
        for r in reader:
            # Datumsfilter (best effort - wenn kein Datum erkennbar, alles zaehlen)
            if date_col and r.get(date_col):
                try:
                    d = parse_iso(str(r[date_col])[:25])
                    if d.tzinfo is None:
                        d = d.replace(tzinfo=von.tzinfo or dt.timezone.utc)
                    if not (von <= d <= bis):
                        continue
                except Exception:
                    pass
            raw = str(r.get(amount_col, "")).replace("'", "").replace(",", ".").strip()
            try:
                total += float(raw)
                count += 1
            except ValueError:
                continue
    return {"total": round(total, 2), "count": count,
            "note": "Datumsfilter aktiv" if date_col else "ohne Datumsfilter (ganze Datei)"}


# ----------------------------------------------------------------------------
# SumUp API
# ----------------------------------------------------------------------------
def _session(api_key):
    import requests
    s = requests.Session()
    s.headers.update({"Authorization": "Bearer " + api_key,
                      "Accept": "application/json"})
    return s


def get_merchant_code(s, cfg):
    if cfg.get("merchant_code"):
        return cfg["merchant_code"]
    r = s.get(EP_ME, timeout=30)
    r.raise_for_status()
    data = r.json()
    return (data.get("merchant_profile") or {}).get("merchant_code") \
        or data.get("merchant_code")


def fetch_transactions(s, von, bis):
    """Alle Transaktionen im Zeitfenster holen (mit Pagination)."""
    items, url = [], EP_TX_HISTORY
    params = {"limit": 100, "order": "ascending",
              "oldest_time": von.isoformat(), "newest_time": bis.isoformat()}
    seen = 0
    while url:
        r = s.get(url, params=params if url == EP_TX_HISTORY else None, timeout=30)
        r.raise_for_status()
        data = r.json()
        batch = data.get("items", data if isinstance(data, list) else [])
        items.extend(batch)
        seen += len(batch)
        # naechste Seite via links[rel=next]
        nxt = None
        for link in data.get("links", []) if isinstance(data, dict) else []:
            if link.get("rel") == "next":
                href = link.get("href", "")
                nxt = href if href.startswith("http") else API_BASE + href
        url, params = nxt, None
        if seen > 50000:  # Sicherheitsbremse
            break
    return items


def fetch_receipt_items(s, mid, tx):
    """Einzelpositionen (Artikel) einer Transaktion holen."""
    tx_id = tx.get("transaction_id") or tx.get("id")
    if not tx_id:
        return []
    r = s.get(EP_RECEIPT.format(tx_id=tx_id, mid=mid), timeout=30)
    if r.status_code != 200:
        return []
    td = (r.json().get("transaction_data") or {})
    out = []
    for p in td.get("products", []) or []:
        out.append({
            "name": p.get("name") or p.get("description") or "Unbenannt",
            "quantity": float(p.get("quantity") or 0),
            "price": float(p.get("price") or 0),
            "total_price": float(p.get("total_price")
                                 or (float(p.get("price") or 0) * float(p.get("quantity") or 0))),
        })
    return out


# ----------------------------------------------------------------------------
# Demodaten (fuer --demo, damit du den Output ohne Keys siehst)
# ----------------------------------------------------------------------------
def demo_data():
    tx = [
        {"transaction_id": "t1", "amount": 35.0, "status": "SUCCESSFUL", "type": "PAYMENT", "payment_type": "POS"},
        {"transaction_id": "t2", "amount": 12.0, "status": "SUCCESSFUL", "type": "PAYMENT", "payment_type": "CASH"},
        {"transaction_id": "t3", "amount": 24.0, "status": "SUCCESSFUL", "type": "PAYMENT", "payment_type": "CASH"},
        {"transaction_id": "t4", "amount": 0.0,  "status": "SUCCESSFUL", "type": "PAYMENT", "payment_type": "CASH"},
        {"transaction_id": "t5", "amount": 18.0, "status": "SUCCESSFUL", "type": "PAYMENT", "payment_type": "POS"},
    ]
    receipts = {
        "t1": [{"name": "Bier 0.5", "quantity": 5, "price": 5.0, "total_price": 25.0},
               {"name": "Vodka Shot 4cl", "quantity": 2, "price": 5.0, "total_price": 10.0}],
        "t2": [{"name": "Bier 0.5", "quantity": 2, "price": 5.0, "total_price": 10.0},
               {"name": "Mineral", "quantity": 1, "price": 2.0, "total_price": 2.0}],
        "t3": [{"name": "Aperol Spritz", "quantity": 3, "price": 8.0, "total_price": 24.0}],
        "t4": [{"name": "Bier 0.5 (Personal)", "quantity": 2, "price": 0.0, "total_price": 0.0}],
        "t5": [{"name": "Bier 0.5", "quantity": 2, "price": 5.0, "total_price": 10.0},
               {"name": "Vodka Shot 4cl", "quantity": 1, "price": 5.0, "total_price": 5.0},
               {"name": "Chips", "quantity": 1, "price": 3.0, "total_price": 3.0}],
    }
    # Demo-EK-Liste
    ek = {norm("Bier 0.5"): ("Bier 0.5", 0.90),
          norm("Bier 0.5 (Personal)"): ("Bier 0.5 (Personal)", 0.90),
          norm("Vodka Shot 4cl"): ("Vodka Shot 4cl", 0.80),
          norm("Aperol Spritz"): ("Aperol Spritz", 2.20),
          norm("Mineral"): ("Mineral", 0.40)}
          # "Chips" fehlt absichtlich -> wird als unbekannt markiert
    twint = {"total": 24.0, "count": 1, "note": "Demo (entspricht t3, bar getippt)"}
    return tx, receipts, ek, twint


# ----------------------------------------------------------------------------
# Kernberechnung
# ----------------------------------------------------------------------------
def compute(transactions, receipts_fn, ek, twint, rate=0.13):
    umsatz = 0.0
    card_total = 0.0
    cash_total = 0.0
    products = {}   # norm -> {"name":..., "qty":..., }
    no_items = 0

    for tx in transactions:
        if tx.get("status") not in (None, "SUCCESSFUL"):
            continue
        ttype = tx.get("type", "PAYMENT")
        amt = float(tx.get("amount") or 0)
        sign = -1.0 if ttype == "REFUND" else 1.0
        umsatz += sign * amt

        ptype = (tx.get("payment_type") or "").upper()
        has_card = bool(tx.get("card_type")) or ptype in CARD_PAYMENT_TYPES
        if ptype in CASH_PAYMENT_TYPES and not tx.get("card_type"):
            cash_total += sign * amt
        elif has_card:
            card_total += sign * amt
        else:
            cash_total += sign * amt  # unbekannt -> zu Bargeld(+TWINT)

        for it in receipts_fn(tx):
            key = norm(it["name"])
            if key not in products:
                products[key] = {"name": it["name"], "qty": 0.0}
            products[key]["qty"] += sign * it["quantity"]
        if not receipts_fn(tx):
            no_items += 1

    # EK je Produkt zuordnen
    rows, unknown = [], []
    getraenkeschuld = 0.0
    for key, p in sorted(products.items(), key=lambda kv: kv[1]["name"].lower()):
        disp, ekval = ek.get(key, (p["name"], None))
        qty = round(p["qty"], 3)
        if ekval is None:
            unknown.append(p["name"])
            line = 0.0
            rows.append({"name": p["name"], "qty": qty, "ek": None, "line": line, "unknown": True})
        else:
            line = round(qty * ekval, 2)
            getraenkeschuld += line
            rows.append({"name": disp, "qty": qty, "ek": ekval, "line": line, "unknown": False})

    umsatz = round(umsatz, 2)
    card_total = round(card_total, 2)
    cash_total = round(cash_total, 2)
    twint_total = round(float(twint.get("total", 0.0)), 2)
    bargeld_real = round(cash_total - twint_total, 2)
    beteiligung = round(umsatz * rate, 2)
    getraenkeschuld = round(getraenkeschuld, 2)
    total_owed = round(beteiligung + getraenkeschuld, 2)

    return {
        "umsatz": umsatz, "rate": rate, "beteiligung": beteiligung,
        "rows": rows, "unknown": unknown, "getraenkeschuld": getraenkeschuld,
        "total_owed": total_owed,
        "card_total": card_total, "cash_total": cash_total,
        "twint_total": twint_total, "bargeld_real": bargeld_real,
        "twint_note": twint.get("note", ""), "no_items": no_items,
        "n_tx": len(transactions),
    }


# ----------------------------------------------------------------------------
# Ausgabe: Excel (mit Formeln) + PDF
# ----------------------------------------------------------------------------
def write_excel(res, meta, path, zahlen=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    FONT = "Arial"
    bold = Font(name=FONT, bold=True)
    base = Font(name=FONT)
    title = Font(name=FONT, bold=True, size=14)
    head_fill = PatternFill("solid", start_color="1F2937")
    head_font = Font(name=FONT, bold=True, color="FFFFFF")
    warn_fill = PatternFill("solid", start_color="FFF3CD")
    money = '#,##0.00;(#,##0.00);"-"'
    thin = Side(style="thin", color="D0D0D0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    wb = Workbook()
    ws = wb.active
    ws.title = "Abrechnung"
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 34
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 16

    def cell(ref, val, f=base, num=None, al=None, fill=None):
        c = ws[ref]; c.value = val; c.font = f
        if num: c.number_format = num
        if al: c.alignment = Alignment(horizontal=al)
        if fill: c.fill = fill
        return c

    cell("A1", "Event-Abrechnung", title)
    cell("A2", meta["event"], bold)
    cell("A3", f"Zeitraum: {meta['von']}  bis  {meta['bis']}", base)
    cell("A4", f"Erstellt: {meta['created']}   Transaktionen: {res['n_tx']}", base)

    r = 6
    cell(f"A{r}", "1) Einnahmen — eingegangen auf dem R1-Konto", bold)
    r += 1
    cell(f"A{r}", "Karte via SumUp (brutto)"); cell(f"D{r}", res["card_total"], base, money, "right")
    card_row = r
    r += 1
    z = zahlen or {}
    cell(f"A{r}", "TWINT (in SumUp als Bargeld getippt, brutto)")
    cell(f"D{r}", z.get("twint_brutto", res["cash_total"]), base, money, "right")
    twintb_row = r
    r += 1
    cell(f"A{r}", "− SumUp-Gebühren (2.5%)")
    cell(f"D{r}", -z.get("sumup_geb", 0), base, money, "right"); sgeb_row = r
    r += 1
    cell(f"A{r}", "− TWINT-Gebühren (1.3%)")
    cell(f"D{r}", -z.get("twint_geb", 0), base, money, "right"); tgeb_row = r
    r += 1
    cell(f"A{r}", "Einnahmen netto", bold)
    cell(f"D{r}", f"=D{card_row}+D{twintb_row}+D{sgeb_row}+D{tgeb_row}", bold, money, "right")
    netto_row = r

    r += 2
    cell(f"A{r}", "2) Abzüge R1 (bleiben beim Trägerverein)", bold)
    r += 1
    cell(f"A{r}", f"Gesamtumsatz (Basis, alle Kategorien)"); cell(f"D{r}", res["umsatz"], base, money, "right")
    umsatz_row = r
    r += 1
    cell(f"A{r}", "Umsatzbeteiligungssatz"); cell(f"D{r}", res["rate"], base, "0.0%", "right")
    rate_row = r
    r += 1
    cell(f"A{r}", "13% Umsatzbeteiligung", bold)
    cell(f"D{r}", f"=D{umsatz_row}*D{rate_row}", bold, money, "right")
    beteiligung_row = r
    r += 1
    cell(f"A{r}", "Einkauf Getränke (Detail unten, inkl. Mitarbeiter/Gratis)", bold)
    ek_row_ref = r  # Formel wird nach der Detailtabelle gesetzt
    r += 1
    cell(f"A{r}", "Miete (nach Wochentag)")
    cell(f"D{r}", z.get("miete", 0), base, money, "right"); miete_row = r
    r += 1
    cell(f"A{r}", "Total Abzüge", bold)
    cell(f"D{r}", f"=D{beteiligung_row}+D{ek_row_ref}+D{miete_row}", bold, money, "right")
    abzuege_row = r

    r += 2
    cell(f"A{r}", "3) ÜBERWEISUNG an die Nutzung (Reingewinn)", title)
    cell(f"D{r}", f"=D{netto_row}-D{abzuege_row}", title, money, "right")
    cell(f"A{r}", "3) ÜBERWEISUNG an die Nutzung (Reingewinn)", title).fill = PatternFill(
        "solid", start_color="DCFCE7")
    ws[f"D{r}"].fill = PatternFill("solid", start_color="DCFCE7")

    r += 2
    cell(f"A{r}", "Detail: Getränke-Einkauf (Menge × Einkaufspreis)", bold)
    r += 1
    for col, txt in zip("ABCD", ["Produkt", "Menge", "EK/Stk", "Total"]):
        c = cell(f"{col}{r}", txt, head_font, fill=head_fill)
        c.alignment = Alignment(horizontal="left" if col == "A" else "right")
        c.border = border
    r += 1
    first = r
    for row in res["rows"]:
        cell(f"A{r}", row["name"], base).border = border
        cell(f"B{r}", row["qty"], base, "#,##0.###", "right").border = border
        if row["unknown"]:
            cell(f"C{r}", "kein EK", base, None, "right").border = border
            cell(f"D{r}", 0, base, money, "right").border = border
        else:
            cell(f"C{r}", row["ek"], base, money, "right").border = border
            cell(f"D{r}", f"=B{r}*C{r}", base, money, "right").border = border
        r += 1
    last = r - 1
    cell(f"A{r}", "Einkauf Getränke (Summe)", bold)
    cell(f"D{r}", f"=SUM(D{first}:D{last})" if last >= first else 0, bold, money, "right")
    # Abzugszeile oben mit der Detailsumme verknuepfen
    cell(f"D{ek_row_ref}", f"=D{r}", bold, money, "right")



    r += 2
    if res["unknown"]:
        cell(f"A{r}", "Ohne EK-Abzug (Fremdkategorie oder nicht in EK-Liste):", bold)
        r += 1
        cell(f"A{r}", ", ".join(res["unknown"]), base)
        r += 1
    if res["no_items"]:
        cell(f"A{r}", f"⚠ {res['no_items']} Transaktion(en) ohne Artikel "
                      f"(als freier Betrag getippt?) — nicht in EK enthalten.", base).fill = warn_fill
        r += 1

    wb.save(path)
    return path


def recalc_excel(path):
    """Formeln in echte Werte rechnen (LibreOffice), falls vorhanden — sonst egal."""
    try:
        import subprocess
        skill_recalc = "/mnt/skills/public/xlsx/scripts/recalc.py"
        if os.path.exists(skill_recalc):
            subprocess.run([sys.executable, skill_recalc, path], timeout=120,
                           capture_output=True)
    except Exception:
        pass


def write_pdf(res, meta, path, zahlen=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    styles = getSampleStyleSheet()
    h = ParagraphStyle("h", parent=styles["Heading1"], fontSize=16, spaceAfter=2)
    sub = ParagraphStyle("s", parent=styles["Normal"], textColor=colors.grey)
    sec = ParagraphStyle("sec", parent=styles["Heading2"], fontSize=11, spaceBefore=10)
    doc = SimpleDocTemplate(path, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm)
    el = [Paragraph("Event-Abrechnung", h),
          Paragraph(meta["event"], styles["Heading2"]),
          Paragraph(f"Zeitraum: {meta['von']} – {meta['bis']}", sub),
          Paragraph(f"Erstellt: {meta['created']} · Transaktionen: {res['n_tx']}", sub),
          Spacer(1, 8)]

    def chf(x): return f"{x:,.2f}".replace(",", "'")

    z = zahlen or {}
    summary = [["Einnahmen R1-Konto: Karte (brutto)", chf(res["card_total"])],
               ["Einnahmen R1-Konto: TWINT, als Bargeld getippt (brutto)",
                chf(z.get("twint_brutto", res["cash_total"]))],
               ["− SumUp-Gebühren (2.5%)", chf(-z.get("sumup_geb", 0))],
               ["− TWINT-Gebühren (1.3%)", chf(-z.get("twint_geb", 0))],
               ["Einnahmen netto", chf(z.get("einnahmen", res["umsatz"]))],
               [f"\u2212 13% Umsatzbeteiligung (auf Umsatz {chf(res['umsatz'])})",
                chf(-res["beteiligung"])],
               ["\u2212 Einkauf Getr\u00e4nke (inkl. Mitarbeiter/Gratis)",
                chf(-res["getraenkeschuld"])],
               ["\u2212 Miete (nach Wochentag)", chf(-z.get("miete", 0))],
               ["\u00dcBERWEISUNG an die Nutzung (Reingewinn)",
                chf(z.get("auszahlung", 0))]]
    t = Table(summary, colWidths=[110*mm, 40*mm])
    t.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), "Helvetica", 10),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
        ("LINEBELOW", (0, 0), (-1, -2), 0.3, colors.lightgrey),
        ("FONT", (0, -1), (-1, -1), "Helvetica-Bold", 11),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#DCFCE7")),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    el += [t, Spacer(1, 6), Paragraph("Detail: Getränke-Einkauf (Menge × Einkaufspreis)", sec)]

    data = [["Produkt", "Menge", "EK/Stk", "Total"]]
    for row in res["rows"]:
        data.append([row["name"], f"{row['qty']:g}",
                     "kein EK" if row["unknown"] else chf(row["ek"]),
                     chf(row["line"])])
    t2 = Table(data, colWidths=[90*mm, 20*mm, 20*mm, 20*mm])
    t2.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), "Helvetica", 9),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F7F7")]),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E5E7EB"))]))
    el += [t2]

    notes = []
    if res["unknown"]:
        notes.append("Ohne EK-Abzug (Fremdkategorie oder nicht in EK-Liste): "
                     + ", ".join(res["unknown"]))
    if res["no_items"]:
        notes.append(f"{res['no_items']} Transaktion(en) ohne Artikel — evtl. als freier "
                     "Betrag getippt; nicht in der EK-Schuld enthalten.")
    if notes:
        el += [Spacer(1, 8)]
        for n in notes:
            el.append(Paragraph("⚠ " + n, ParagraphStyle(
                "warn", parent=styles["Normal"], textColor=colors.HexColor("#92400E"))))
    doc.build(el)
    return path


# ----------------------------------------------------------------------------
# Optional: KI-Schritt (Abrechnungs-E-Mail formulieren). Standardmaessig AUS.
# Benoetigt ANTHROPIC_API_KEY in der Umgebung.
# ----------------------------------------------------------------------------
def ki_email_text(res, meta):
    import requests
    key = os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        return None
    prompt = (
        "Schreibe eine kurze, freundliche Abrechnungs-E-Mail auf Deutsch (Du-Form, "
        "schweizerischer Ton) an eine Nutzung unseres Kulturraums. Zahlen:\n"
        f"- Event: {meta['event']}\n- Gesamtumsatz: CHF {res['umsatz']:.2f}\n"
        f"- Umsatzbeteiligung {res['rate']*100:.0f}%: CHF {res['beteiligung']:.2f}\n"
        f"- Getränkeschuld (Einkaufspreis): CHF {res['getraenkeschuld']:.2f}\n"
        f"- TOTAL offen: CHF {res['total_owed']:.2f}\n"
        "Kontext: Das Geld ist bereits beim Trägerverein eingegangen; wir überweisen "
        "der Nutzung ihren Reingewinn. Maximal 8 Sätze, keine Floskeln.")
    r = requests.post("https://api.anthropic.com/v1/messages",
                      headers={"x-api-key": key, "anthropic-version": "2023-06-01",
                               "content-type": "application/json"},
                      json={"model": "claude-sonnet-4-6", "max_tokens": 500,
                            "messages": [{"role": "user", "content": prompt}]},
                      timeout=60)
    r.raise_for_status()
    return "".join(b.get("text", "") for b in r.json().get("content", []))


# ----------------------------------------------------------------------------
# EVENTBLATT R1: Vorlage mit Tokens bauen + Tokens befuellen
#
# Es werden NUR finanzielle Felder befuellt. Nicht-finanzielle Felder (Eventangaben,
# Tickets, Unterschriften, sonstige Kosten) bleiben leer fuer die manuelle Eingabe.
# ----------------------------------------------------------------------------
def chf(x):
    if x is None or x == "":
        return ""
    return f"{float(x):,.2f}".replace(",", "'")


def build_eventblatt_vorlage(path):
    """Erstellt eine EVENTBLATT-Vorlage mit {{TOKENS}} in den Finanzfeldern."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    base = d.styles["Normal"]
    base.font.name = "Arial"
    base.font.size = Pt(10)

    def h(text, size=14):
        p = d.add_paragraph()
        r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
        return p

    def it(text):
        p = d.add_paragraph()
        r = p.add_run(text); r.italic = True; r.font.size = Pt(8)
        return p

    def kv(rows):
        t = d.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        t.columns[0].width = Pt(260); t.columns[1].width = Pt(120)
        for label, token in rows:
            c = t.add_row().cells
            rp = c[0].paragraphs[0]; rr = rp.add_run(label); rr.bold = True
            if token:
                c[1].paragraphs[0].add_run(token)
        return t

    h("EVENTBLATT", 16)
    h("Abrechnung Getränke und Ticketeinnahmen", 11)
    it("Beilage zum Kooperationsvertrag – von verantwortlicher Person zu unterzeichnen")

    h("Eventangaben", 12)
    it("Erster Teil direkt nach dem Event und noch vor Ort ausfüllen und unterzeichnen!")
    kv([("Event:", ""), ("Datum:", ""),
        ("Verantwortliche Person Kooperationspartner:", ""),
        ("Tel. Verantwortliche Person Kooperationspartner:", ""),
        ("Abendverantwortung R1:", "")])
    d.add_paragraph()
    kv([("Bruttoeinnahmen SumUp (CHF):", ""),
        ("Anzahl Abendkasse Eintritte (CHF):", ""),
        ("Abendkasse Kosten pro Eintritt (CHF):", "")])
    it("Die unterzeichnende verantwortliche Person bestätigt die Richtigkeit der obenstehenden Angaben.")
    d.add_paragraph("Ort, Datum: ______________________   Unterschrift Kooperationspartner: ______________________")
    d.add_paragraph("Ort, Datum: ______________________   Unterschrift Abendverantwortung: ______________________")

    d.add_paragraph("―" * 40)
    h("ABRECHNUNG", 14)
    d.add_paragraph("Abrechnung durchgeführt von: ______________________________")

    h("Einnahmen SumUp", 12)
    kv([("Total Einnahmen Brutto (CHF):", "{{SUMUP_BRUTTO}}"),
        ("Gebühren (CHF):", "{{SUMUP_GEBUEHR}}"),
        ("Total Einnahmen Netto (CHF):", "{{SUMUP_NETTO}}")])

    h("Einnahmen Twint", 12)
    kv([("Total Einnahmen Brutto (CHF):", "{{TWINT_BRUTTO}}"),
        ("Gebühren (CHF):", "{{TWINT_GEBUEHR}}"),
        ("Total Einnahmen Netto (CHF):", "{{TWINT_NETTO}}")])

    h("Einnahmen Tickets", 12)
    it("Manuell – externe Ticketing-Plattform.")
    kv([("Ticketing-Plattform:", ""), ("Anzahl verkaufte Tickets:", ""),
        ("Kosten pro Ticket (CHF):", ""), ("Total Einnahmen Brutto (CHF):", ""),
        ("Gebühren (CHF):", ""), ("Total Einnahmen Netto (CHF):", ""),
        ("Nachweis beigefügt (Rapport/Screenshot):", "")])

    h("Abzüge", 12)
    kv([("13% Umsatzbeteiligung (CHF):", "{{UMSATZBETEILIGUNG}}"),
        ("Tagespauschale (CHF):", "{{TAGESPAUSCHALE}}"),
        ("Reinigungskosten (CHF):", "{{REINIGUNG}}"),
        ("Einkaufskosten Getränke (CHF):", "{{EK_GETRAENKE}}"),
        ("Abendverantwortung (CHF):", "{{ABENDVERANTWORTUNG}}"),
        ("Abzug sonstige Kosten (CHF):", ""),
        ("Total Abzüge (CHF):", "{{TOTAL_ABZUEGE}}")])

    h("Gesamtabrechnung", 12)
    kv([("Total Einnahmen (CHF):", "{{TOTAL_EINNAHMEN}}"),
        ("Total Abzüge (CHF):", "{{TOTAL_ABZUEGE}}"),
        ("Auszahlung an Kooperationspartner (CHF):", "{{AUSZAHLUNG}}")])
    it("Die unterzeichnende verantwortliche Person bestätigt die Richtigkeit der obenstehenden Angaben.")
    d.add_paragraph("Ort, Datum: ______________________   Unterschrift Kooperationspartner: ______________________")
    d.add_paragraph("Ort, Datum: ______________________   Unterschrift Abrechnungsverantwortung Trägerverein: ______________________")

    d.save(path)
    return path


def _replace_tokens_in_paragraph(p, mapping):
    """Token in einem Absatz ersetzen (auch wenn ueber mehrere Runs verteilt)."""
    full = "".join(run.text for run in p.runs)
    if "{{" not in full:
        return
    new = full
    for k, v in mapping.items():
        new = new.replace(k, v)
    if new != full and p.runs:
        p.runs[0].text = new
        for run in p.runs[1:]:
            run.text = ""


def fill_eventblatt(template_path, output_path, mapping):
    """Befuellt alle {{TOKENS}} in einem docx (Absaetze + Tabellenzellen)."""
    from docx import Document
    d = Document(template_path)

    def walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_tokens_in_paragraph(p, mapping)
                    walk_tables(cell.tables)

    for p in d.paragraphs:
        _replace_tokens_in_paragraph(p, mapping)
    walk_tables(d.tables)
    d.save(output_path)
    return output_path


def miete_fuer_datum(d):
    """Raummiete nach Wochentag des Eventstarts: So-Mi 50, Do 100, Fr+Sa 200."""
    wd = d.weekday()  # Mo=0 ... So=6
    if wd in (6, 0, 1, 2):   # So, Mo, Di, Mi
        return 50.0
    if wd == 3:              # Do
        return 100.0
    return 200.0             # Fr, Sa


def eventblatt_values(res, cfg, ticket_netto=0.0, event_date=None):
    """Berechnet die Finanzwerte fuers Eventblatt aus dem Abrechnungs-Ergebnis.

    Gebuehren-Logik (Schaetzung, konfigurierbar):
      - SumUp-Gebuehr auf alles, was mit Karte getippt wurde (Default 2.5%).
      - TWINT-Gebuehr auf alles, was als 'Bargeld' getippt wurde (Default 1.3%),
        denn im R1-Workflow steht 'bar getippt' fuer TWINT-QR-Zahlungen.
        Liegt eine RaiseNow-CSV vor, wird stattdessen deren Total verwendet.
      - Miete automatisch nach Wochentag (So-Mi 50 / Do 100 / Fr+Sa 200),
        via config['eventblatt']['tagespauschale'] uebersteuerbar.
    """
    eb = (cfg.get("eventblatt") or {})
    sumup_rate = float(eb.get("sumup_fee_rate", 0.025))
    twint_rate = float(eb.get("twint_fee_rate", 0.013))

    sumup_brutto = res["card_total"]                 # Karte via SumUp
    twint_brutto = res["cash_total"]                 # bar getippt = TWINT (kein Bargeld moeglich)
    sumup_geb = round(sumup_brutto * sumup_rate, 2)
    twint_geb = round(twint_brutto * twint_rate, 2)
    sumup_netto = round(sumup_brutto - sumup_geb, 2)
    twint_netto = round(twint_brutto - twint_geb, 2)

    umsatzbeteiligung = res["beteiligung"]
    ek = res["getraenkeschuld"]
    tagespauschale = eb.get("tagespauschale")
    if tagespauschale in (None, "") and event_date is not None:
        tagespauschale = miete_fuer_datum(event_date)
    reinigung = eb.get("reinigung")
    abendverantwortung = eb.get("abendverantwortung")

    abzuege = umsatzbeteiligung + ek
    for v in (tagespauschale, reinigung, abendverantwortung):
        if v not in (None, ""):
            abzuege += float(v)
    abzuege = round(abzuege, 2)
    ticket_netto = float(ticket_netto or 0.0)
    einnahmen = round(sumup_netto + twint_netto + ticket_netto, 2)
    auszahlung = round(einnahmen - abzuege, 2)

    def opt(v):  # optionaler Pauschalwert -> Zahl oder leer
        return chf(v) if v not in (None, "") else ""

    mapping = {
        "{{SUMUP_BRUTTO}}": chf(sumup_brutto), "{{SUMUP_GEBUEHR}}": chf(sumup_geb),
        "{{SUMUP_NETTO}}": chf(sumup_netto),
        "{{TWINT_BRUTTO}}": chf(twint_brutto), "{{TWINT_GEBUEHR}}": chf(twint_geb),
        "{{TWINT_NETTO}}": chf(twint_netto),
        "{{UMSATZBETEILIGUNG}}": chf(umsatzbeteiligung),
        "{{EK_GETRAENKE}}": chf(ek),
        "{{TAGESPAUSCHALE}}": opt(tagespauschale), "{{REINIGUNG}}": opt(reinigung),
        "{{ABENDVERANTWORTUNG}}": opt(abendverantwortung),
        "{{TOTAL_ABZUEGE}}": chf(abzuege), "{{TOTAL_EINNAHMEN}}": chf(einnahmen),
        "{{AUSZAHLUNG}}": chf(auszahlung),
    }
    zahlen = {"sumup_brutto": sumup_brutto, "sumup_geb": sumup_geb,
              "sumup_netto": sumup_netto, "twint_brutto": twint_brutto,
              "twint_geb": twint_geb, "twint_netto": twint_netto,
              "umsatzbeteiligung": umsatzbeteiligung, "ek": ek,
              "miete": float(tagespauschale or 0), "abzuege": abzuege,
              "einnahmen": einnahmen, "auszahlung": auszahlung,
              "ticket_netto": ticket_netto}
    return mapping, zahlen


# ----------------------------------------------------------------------------
# AUTOMATIK: Event-Erkennung (4h Stille = Event zu Ende), Abrechnung, Mailversand
# ----------------------------------------------------------------------------
def tx_time(t):
    for k in ("timestamp", "time", "created_at", "local_time"):
        if t.get(k):
            try:
                return parse_iso(str(t[k]))
            except Exception:
                continue
    return None


def cluster_events(transactions, gap_hours=4.0):
    """Sortiert Transaktionen zeitlich und trennt Events an Luecken > gap_hours."""
    ts = [(tx_time(t), t) for t in transactions]
    ts = sorted([x for x in ts if x[0] is not None], key=lambda x: x[0])
    events, cur = [], []
    for when, t in ts:
        if cur and (when - cur[-1][0]).total_seconds() > gap_hours * 3600:
            events.append(cur)
            cur = []
        cur.append((when, t))
    if cur:
        events.append(cur)
    return events  # Liste von Listen [(zeit, tx), ...]


def load_state(path):
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"abgerechnet": []}


def save_state(path, state):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(state, f, indent=2)


def send_mail(cfg, subject, body, attachments):
    """Versendet die Abrechnung per SMTP (z.B. Gmail mit App-Passwort)."""
    import smtplib
    from email.message import EmailMessage
    import mimetypes

    em = (cfg.get("email") or {})
    host = em.get("smtp_host", "smtp.gmail.com")
    port = int(em.get("smtp_port", 587))
    user = em.get("user") or os.environ.get("MAIL_USER")
    pw = em.get("app_password") or os.environ.get("MAIL_APP_PASSWORD")
    to = em.get("to") or os.environ.get("MAIL_TO")
    if not (user and pw and to):
        raise RuntimeError("E-Mail nicht konfiguriert: config.json -> email "
                           "{smtp_host, smtp_port, user, app_password, to}")

    msg = EmailMessage()
    msg["From"] = user
    msg["To"] = to
    msg["Subject"] = subject
    msg.set_content(body)
    for path in attachments:
        if not path or not os.path.exists(path):
            continue
        ctype, _ = mimetypes.guess_type(path)
        maintype, subtype = (ctype or "application/octet-stream").split("/", 1)
        with open(path, "rb") as f:
            msg.add_attachment(f.read(), maintype=maintype, subtype=subtype,
                               filename=os.path.basename(path))
    with smtplib.SMTP(host, port, timeout=60) as srv:
        srv.starttls()
        srv.login(user, pw)
        srv.send_message(msg)


def run_event(tx, receipts_fn, ek, twint, cfg, event_name, von_s, bis_s,
              out_dir, rate, event_date=None):
    """Kompletter Abrechnungslauf fuer ein Event -> (res, zahlen, pfade)."""
    res = compute(tx, receipts_fn, ek, twint, rate=rate)
    mapping, zahlen = eventblatt_values(res, cfg, event_date=event_date)
    created = dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    meta = {"event": event_name, "von": von_s, "bis": bis_s, "created": created}
    safe = "".join(c if c.isalnum() or c in "-_ " else "_"
                   for c in event_name).strip().replace(" ", "_")
    os.makedirs(out_dir, exist_ok=True)
    xlsx_path = os.path.join(out_dir, f"Abrechnung_{safe}.xlsx")
    pdf_path = os.path.join(out_dir, f"Abrechnung_{safe}.pdf")
    write_excel(res, meta, xlsx_path, zahlen)
    recalc_excel(xlsx_path)
    write_pdf(res, meta, pdf_path, zahlen)

    eventblatt_path = None
    try:
        eb_cfg = (cfg.get("eventblatt") or {})
        template = eb_cfg.get("template")
        if not template or not os.path.exists(template):
            template = os.path.join(out_dir, "EVENTBLATT_Vorlage.docx")
            if not os.path.exists(template):
                build_eventblatt_vorlage(template)
        eventblatt_path = os.path.join(out_dir, f"Eventblatt_{safe}.docx")
        fill_eventblatt(template, eventblatt_path, mapping)
    except Exception as e:
        print(f"  ⚠ Eventblatt konnte nicht erstellt werden: {e}")

    return res, zahlen, {"xlsx": xlsx_path, "pdf": pdf_path,
                         "eventblatt": eventblatt_path}


def mail_body(event_name, res, zahlen):
    z = zahlen or {}
    lines = [
        f"Automatische Abrechnung: {event_name}",
        "",
        f"Gesamtumsatz (alle Kategorien):  CHF {res['umsatz']:.2f}",
        f"  davon Karte (SumUp):           CHF {res['card_total']:.2f}",
        f"  davon TWINT (bar getippt):     CHF {res['cash_total']:.2f}",
        "",
        "Abzüge:",
        f"  13% Umsatzbeteiligung:         CHF {res['beteiligung']:.2f}",
        f"  Einkauf Getränke (inkl. Mitarbeiter/Gratis): CHF {res['getraenkeschuld']:.2f}",
    ]
    if z:
        lines += [
            f"  Miete (nach Wochentag):        CHF {z['miete']:.2f}",
            f"  SumUp-Gebühren (2.5% auf Karte): CHF {z['sumup_geb']:.2f}",
            f"  TWINT-Gebühren (1.3% auf TWINT): CHF {z['twint_geb']:.2f}",
            "",
            f"Einnahmen netto:                 CHF {z['einnahmen']:.2f}",
            f"Total Abzüge:                    CHF {z['abzuege']:.2f}",
            f"ÜBERWEISUNG an euch (Reingewinn): CHF {z['auszahlung']:.2f}",
        ]
    if res.get("unknown"):
        lines += ["", "Ohne EK-Abzug (Fremdkategorie): " + ", ".join(res["unknown"])]
    if res.get("no_items"):
        lines += [f"Hinweis: {res['no_items']} Transaktion(en) ohne Artikel "
                  f"(freier Betrag getippt) – ohne EK-Abzug."]
    lines += ["", "Details im angehängten Eventblatt (Word), Excel und PDF.",
              "Gebührensätze: SumUp 2.5% auf Kartenzahlungen, TWINT 1.3%."]
    return "\n".join(lines)


def auto_run(args, cfg, once=True):
    """Prueft auf beendete Events (>=4h Stille) und rechnet neue automatisch ab."""
    api_key = cfg.get("sumup_api_key") or os.environ.get("SUMUP_API_KEY")
    if not api_key:
        raise SystemExit("SumUp API-Key fehlt (config.json -> sumup_api_key "
                         "oder Umgebungsvariable SUMUP_API_KEY).")
    s = _session(api_key)
    mid = get_merchant_code(s, cfg)
    ek = load_ek_prices(args.ek)
    gap = float((cfg.get("auto") or {}).get("gap_hours", 4))
    lookback = float((cfg.get("auto") or {}).get("lookback_hours", args.lookback_hours))
    state_path = os.path.join(args.out, "auto_state.json")

    while True:
        now = dt.datetime.now(dt.timezone.utc)
        von = now - dt.timedelta(hours=lookback)
        try:
            tx = fetch_transactions(s, von, now)
        except Exception as e:
            print(f"[{now:%H:%M}] API-Fehler: {e}")
            tx = []
        state = load_state(state_path)
        done_ids = set(state.get("abgerechnet", []))
        events = cluster_events(tx, gap_hours=gap)
        neu = 0
        for ev in events:
            first_t, last_t = ev[0][0], ev[-1][0]
            ev_id = last_t.isoformat()
            fertig = (now - last_t).total_seconds() >= gap * 3600
            if not fertig or ev_id in done_ids:
                continue
            lokal = first_t.astimezone()
            name = f"Event_{lokal:%Y-%m-%d}"
            print(f"→ Abgeschlossenes Event erkannt: {name} "
                  f"({first_t:%d.%m %H:%M} – {last_t:%d.%m %H:%M} UTC, "
                  f"{len(ev)} Transaktionen)")
            ev_tx = [t for _, t in ev]
            twint = {"total": 0.0, "count": 0,
                     "note": "Auto-Modus: bar getippt = TWINT (Gebührenbasis)"}
            receipts_fn = lambda t: fetch_receipt_items(s, mid, t)
            res, zahlen, paths = run_event(
                ev_tx, receipts_fn, ek, twint, cfg, name,
                first_t.astimezone().isoformat(), last_t.astimezone().isoformat(),
                args.out, args.rate, event_date=lokal)
            try:
                send_mail(cfg, f"R1 Abrechnung {name} – Auszahlung CHF "
                               f"{(zahlen or {}).get('auszahlung', 0):.2f}",
                          mail_body(name, res, zahlen),
                          [paths["eventblatt"], paths["xlsx"], paths["pdf"]])
                print(f"  ✉ Mail versendet an {(cfg.get('email') or {}).get('to')}")
            except Exception as e:
                print(f"  ⚠ Mailversand fehlgeschlagen: {e} "
                      f"(Dateien liegen in {args.out})")
            done_ids.add(ev_id)
            state["abgerechnet"] = sorted(done_ids)[-200:]
            save_state(state_path, state)
            neu += 1
        if neu == 0:
            print(f"[{dt.datetime.now():%d.%m %H:%M}] Kein neues beendetes Event.")
        if once:
            return
        import time
        time.sleep(15 * 60)


# ----------------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(description="Tutti Tschutti Event-Abrechnung (Variante 3)")
    ap.add_argument("--event", default="Event")
    ap.add_argument("--von", help="Start ISO-8601, z.B. 2026-06-11T17:00:00+02:00")
    ap.add_argument("--bis", help="Ende ISO-8601")
    ap.add_argument("--twint", help="Pfad zur RaiseNow/TWINT-CSV")
    ap.add_argument("--ek", default="ek_preise.csv", help="Pfad zur EK-Preisliste (CSV)")
    ap.add_argument("--config", default="config.json", help="Pfad zur config.json")
    ap.add_argument("--rate", type=float, default=0.13, help="Umsatzbeteiligung (Default 0.13)")
    ap.add_argument("--out", default=".", help="Ausgabeordner")
    ap.add_argument("--demo", action="store_true", help="Mit eingebauten Demodaten testen")
    ap.add_argument("--inspect", action="store_true",
                    help="Nur die vorkommenden payment_type/card_type Werte ausgeben")
    ap.add_argument("--ki-email", action="store_true", help="Abrechnungs-E-Mail via KI erzeugen")
    ap.add_argument("--test-connection", action="store_true",
                    help="Prüft nur, ob der SumUp-API-Key funktioniert (ruft /me auf).")
    ap.add_argument("--auto", action="store_true",
                    help="Automatik: beendete Events (>=4h Stille) abrechnen + mailen, dann Ende. Für cron.")
    ap.add_argument("--watch", action="store_true",
                    help="Wie --auto, aber als Dauerschleife (prüft alle 15 Min).")
    ap.add_argument("--lookback-hours", type=float, default=72,
                    help="Automatik: wie weit zurück nach Events gesucht wird (Default 72h).")
    args = ap.parse_args()

    cfg = load_config(args.config)

    if args.test_connection:
        api_key = cfg.get("sumup_api_key") or os.environ.get("SUMUP_API_KEY")
        if not api_key:
            ap.error("SumUp API-Key fehlt (config.json -> sumup_api_key oder Umgebungsvariable SUMUP_API_KEY).")
        s = _session(api_key)
        r = s.get(EP_ME, timeout=30)
        if r.status_code != 200:
            print(f"❌ Verbindung fehlgeschlagen (HTTP {r.status_code}): {r.text[:200]}")
            print("   -> Key falsch/abgelaufen oder ohne Leserechte?")
            return
        data = r.json()
        prof = data.get("merchant_profile") or {}
        print("✅ Verbindung OK")
        print("   Händler:", prof.get("company_name") or prof.get("legal_name") or "(unbekannt)")
        print("   merchant_code:", prof.get("merchant_code") or data.get("merchant_code") or "(nicht gefunden)")
        return

    if args.auto or args.watch:
        auto_run(args, cfg, once=args.auto and not args.watch)
        return

    if args.demo:
        tx, receipts, ek, twint = demo_data()
        receipts_fn = lambda t: receipts.get(t.get("transaction_id"), [])
        von_s, bis_s = "Demo", "Demo"
        event_date = dt.datetime.now()
    else:
        if not (args.von and args.bis):
            ap.error("--von und --bis sind im Echtlauf nötig (oder nutze --demo).")
        von, bis = parse_iso(args.von), parse_iso(args.bis)
        von_s, bis_s = args.von, args.bis
        api_key = cfg.get("sumup_api_key") or os.environ.get("SUMUP_API_KEY")
        if not api_key:
            ap.error("SumUp API-Key fehlt (config.json -> sumup_api_key oder Umgebungsvariable SUMUP_API_KEY).")
        s = _session(api_key)
        mid = get_merchant_code(s, cfg)
        tx = fetch_transactions(s, von, bis)

        if args.inspect:
            seen = {}
            for t in tx:
                k = (t.get("payment_type"), t.get("card_type"), t.get("type"))
                seen[k] = seen.get(k, 0) + 1
            print("Vorkommende (payment_type, card_type, type) Werte:")
            for k, n in sorted(seen.items(), key=lambda x: -x[1]):
                print(f"  {k}: {n}x")
            print("\n-> CARD_PAYMENT_TYPES / CASH_PAYMENT_TYPES im Script ggf. anpassen.")
            return

        ek = load_ek_prices(args.ek)
        twint = load_twint_csv(args.twint, von, bis, cfg)
        receipts_fn = lambda t: fetch_receipt_items(s, mid, t)
        event_date = von

    res, zahlen, paths = run_event(tx, receipts_fn, ek, twint, cfg, args.event,
                                   von_s, bis_s, args.out, args.rate,
                                   event_date=event_date)
    xlsx_path, pdf_path = paths["xlsx"], paths["pdf"]
    eventblatt_path = paths["eventblatt"]
    safe = "".join(c if c.isalnum() or c in "-_ " else "_"
                   for c in args.event).strip().replace(" ", "_")
    created = dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    meta = {"event": args.event, "von": von_s, "bis": bis_s, "created": created}

    print(f"\n=== {args.event} ===")
    print(f"Umsatz:            CHF {res['umsatz']:.2f}")
    print(f"Beteiligung {args.rate*100:.0f}%:   CHF {res['beteiligung']:.2f}")
    print(f"Getränkeschuld:    CHF {res['getraenkeschuld']:.2f}")
    print(f"Abzüge R1 (13%+EK): CHF {res['total_owed']:.2f}")
    print(f"  Karte: {res['card_total']:.2f} | TWINT (bar getippt): {res['cash_total']:.2f}")
    if res["unknown"]:
        print("  ohne EK-Abzug (Fremdkategorie):", ", ".join(res["unknown"]))
    if zahlen:
        print(f"  Miete: {zahlen['miete']:.2f} | SumUp-Geb. {zahlen['sumup_geb']:.2f} | "
              f"TWINT-Geb. {zahlen['twint_geb']:.2f}")
        print(f"  ÜBERWEISUNG an Nutzung: CHF {zahlen['auszahlung']:.2f}")
    print(f"\nGeschrieben:\n  {xlsx_path}\n  {pdf_path}")
    if eventblatt_path:
        print(f"  {eventblatt_path}")

    if args.ki_email:
        txt = ki_email_text(res, meta)
        if txt:
            with open(os.path.join(args.out, f"Email_{safe}.txt"), "w", encoding="utf-8") as f:
                f.write(txt)
            print(f"  {os.path.join(args.out, f'Email_{safe}.txt')}")
        else:
            print("  (KI-E-Mail übersprungen: ANTHROPIC_API_KEY nicht gesetzt)")


if __name__ == "__main__":
    main()
